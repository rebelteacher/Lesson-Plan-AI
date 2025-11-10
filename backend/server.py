from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, Request
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
from passlib.context import CryptContext
import jwt
import json
from emergentintegrations.llm.chat import LlmChat, UserMessage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from fastapi.responses import StreamingResponse

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# JWT settings
JWT_SECRET = os.environ.get('JWT_SECRET', 'your_jwt_secret_key_change_in_production')
JWT_ALGORITHM = os.environ.get('JWT_ALGORITHM', 'HS256')
JWT_EXPIRATION_HOURS = int(os.environ.get('JWT_EXPIRATION_HOURS', 24))

# Security
security = HTTPBearer()

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")


# Models
class InvitationCode(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    code: str
    created_by: str  # admin user id
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    used_by: Optional[str] = None  # user id who used it
    used_at: Optional[datetime] = None
    is_active: bool = True

class CreateInvitationCode(BaseModel):
    count: int = 1  # number of codes to generate

class UserRegister(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    state: Optional[str] = None
    invitation_code: str  # Required invitation code

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    full_name: str
    state: Optional[str] = None
    role: str = "teacher"  # teacher or admin
    is_active: bool = True
    join_code: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    last_login: Optional[datetime] = None

class LessonPlanCreate(BaseModel):
    textbook: str
    start_date: str
    end_date: str
    lesson_range: str
    next_major_assessment: str
    state_standards: Optional[str] = None

class DayPlan(BaseModel):
    day_name: str  # e.g., "Monday", "Tuesday"
    day_date: str  # e.g., "2025-01-15"
    learner_outcomes: str
    standards: str
    materials_needed: str
    anticipatory_set: str
    teaching_lesson: str
    modeling: str
    instructional_strategies: str
    check_understanding: str
    guided_practice: str
    independent_practice: str
    closure: str
    summative_assessment: str
    formative_assessment: str
    extended_activities: str
    review_reteach: str
    early_finishers: str

class LessonPlan(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    textbook: str
    start_date: str
    end_date: str
    lesson_range: str
    next_major_assessment: str
    daily_plans: List[DayPlan]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AdminStats(BaseModel):
    total_users: int
    active_users: int
    inactive_users: int
    total_lesson_plans: int

class UserDetail(BaseModel):
    id: str
    email: str
    full_name: str
    state: Optional[str]
    is_active: bool
    join_code: Optional[str]
    created_at: str
    last_login: Optional[str]
    lesson_plan_count: int


# Helper functions
def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)
    return encoded_jwt

def generate_join_code() -> str:
    return str(uuid.uuid4())[:8].upper()

def get_weekdays_between(start_date_str: str, end_date_str: str):
    from datetime import datetime as dt
    start = dt.strptime(start_date_str, "%Y-%m-%d")
    end = dt.strptime(end_date_str, "%Y-%m-%d")
    days = []
    current = start
    while current <= end:
        # 0 = Monday, 6 = Sunday
        if current.weekday() < 5:  # Monday to Friday
            days.append({
                'date': current.strftime("%Y-%m-%d"),
                'day_name': current.strftime("%A")
            })
        current += timedelta(days=1)
    return days

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid authentication credentials")
        
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if user is None:
            raise HTTPException(status_code=401, detail="User not found")
        
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_admin_user(current_user: dict = Depends(get_current_user)):
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user


# Auth routes
@api_router.post("/auth/register")
async def register(user_data: UserRegister):
    # Validate invitation code
    invitation = await db.invitation_codes.find_one({
        "code": user_data.invitation_code,
        "is_active": True
    })
    
    if not invitation:
        raise HTTPException(status_code=400, detail="Invalid or inactive invitation code")
    
    if invitation.get('used_by'):
        raise HTTPException(status_code=400, detail="This invitation code has already been used")
    
    # Check if user exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create user
    user = User(
        email=user_data.email,
        full_name=user_data.full_name,
        state=user_data.state,
        join_code=None  # No longer needed
    )
    
    user_dict = user.model_dump()
    user_dict['password'] = hash_password(user_data.password)
    user_dict['created_at'] = user_dict['created_at'].isoformat()
    user_dict['invitation_code'] = user_data.invitation_code
    
    await db.users.insert_one(user_dict)
    
    # Mark invitation code as used
    await db.invitation_codes.update_one(
        {"code": user_data.invitation_code},
        {
            "$set": {
                "used_by": user.id,
                "used_at": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    # Create token
    token = create_access_token({"sub": user.id, "email": user.email, "role": user.role})
    
    return {
        "token": token,
        "user": {
            "id": user.id,
            "email": user.email,
            "full_name": user.full_name,
            "role": user.role,
            "state": user.state
        }
    }

@api_router.post("/auth/login")
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user['password']):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    if not user.get('is_active', True):
        raise HTTPException(status_code=403, detail="Account is deactivated")
    
    # Update last login
    await db.users.update_one(
        {"id": user['id']},
        {"$set": {"last_login": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Create token
    token = create_access_token({"sub": user['id'], "email": user['email'], "role": user.get('role', 'teacher')})
    
    return {
        "token": token,
        "user": {
            "id": user['id'],
            "email": user['email'],
            "full_name": user['full_name'],
            "role": user.get('role', 'teacher'),
            "join_code": user.get('join_code'),
            "state": user.get('state')
        }
    }

@api_router.get("/auth/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    return {
        "id": current_user['id'],
        "email": current_user['email'],
        "full_name": current_user['full_name'],
        "role": current_user.get('role', 'teacher'),
        "join_code": current_user.get('join_code'),
        "state": current_user.get('state'),
        "is_active": current_user.get('is_active', True)
    }


# Lesson plan routes
@api_router.post("/lesson-plans")
async def create_lesson_plan(plan_data: LessonPlanCreate, current_user: dict = Depends(get_current_user)):
    try:
        # Get weekdays between start and end date
        weekdays = get_weekdays_between(plan_data.start_date, plan_data.end_date)
        
        if not weekdays:
            raise HTTPException(status_code=400, detail="No weekdays found in the date range")
        
        # Initialize Claude chat
        api_key = os.environ.get('EMERGENT_LLM_KEY')
        
        daily_plans = []
        
        # Generate plan for each day
        for idx, day_info in enumerate(weekdays):
            chat = LlmChat(
                api_key=api_key,
                session_id=f"lesson_plan_{current_user['id']}_{day_info['date']}",
                system_message="You are an expert education consultant helping teachers create detailed daily lesson plans."
            )
            chat.with_model("anthropic", "claude-3-7-sonnet-20250219")
            
            state_standards_text = f"\nState Standards to Align With: {plan_data.state_standards}" if plan_data.state_standards else ""
            
            prompt = f"""Create a detailed lesson plan for {day_info['day_name']}, {day_info['date']} (Day {idx+1} of {len(weekdays)}) based on:

Textbook: {plan_data.textbook}
Lesson Range: {plan_data.lesson_range}
Overall Date Range: {plan_data.start_date} to {plan_data.end_date}
Next Major Assessment: {plan_data.next_major_assessment}{state_standards_text}

Provide specific, actionable content for THIS DAY ONLY for each section:

1. Learner Outcomes/Objectives
2. Standards (include the relevant state standards provided above, formatted clearly)
3. Materials Needed
4. Anticipatory Set
5. Teaching the Lesson
6. Modeling
7. Instructional Strategies
8. Check for Understanding
9. Guided Practice/Monitoring
10. Independent Practice
11. Closure
12. Summative Assessment
13. Formative Assessment
14. Extended Activities
15. Review and Reteach Activities
16. Early Finishers Activities

Make each section detailed and specific to day {idx+1}."""
            
            user_message = UserMessage(text=prompt)
            response = await chat.send_message(user_message)
            response_text = response if isinstance(response, str) else str(response)
            
            # Parse AI response into sections
            sections = {
                'learner_outcomes': '',
                'standards': '',
                'materials_needed': '',
                'anticipatory_set': '',
                'teaching_lesson': '',
                'modeling': '',
                'instructional_strategies': '',
                'check_understanding': '',
                'guided_practice': '',
                'independent_practice': '',
                'closure': '',
                'summative_assessment': '',
                'formative_assessment': '',
                'extended_activities': '',
                'review_reteach': '',
                'early_finishers': ''
            }
            
            # Simple section parsing
            lines = response_text.split('\n')
            current_section = None
            current_text = []
            
            section_mapping = {
                'learner outcomes': 'learner_outcomes',
                'objectives': 'learner_outcomes',
                'standards': 'standards',
                'materials needed': 'materials_needed',
                'anticipatory set': 'anticipatory_set',
                'teaching the lesson': 'teaching_lesson',
                'modeling': 'modeling',
                'instructional strategies': 'instructional_strategies',
                'check for understanding': 'check_understanding',
                'guided practice': 'guided_practice',
                'monitoring': 'guided_practice',
                'independent practice': 'independent_practice',
                'closure': 'closure',
                'summative assessment': 'summative_assessment',
                'formative assessment': 'formative_assessment',
                'extended activities': 'extended_activities',
                'review and reteach': 'review_reteach',
                'reteach activities': 'review_reteach',
                'early finishers': 'early_finishers'
            }
            
            for line in lines:
                line_lower = line.lower().strip()
                found_section = False
                
                # Check if this line is a section header
                for keyword, section_key in section_mapping.items():
                    if keyword in line_lower and (line.startswith('#') or line.startswith('**') or line.endswith(':') or len(line) < 50):
                        if current_section and current_text:
                            sections[current_section] = '\n'.join(current_text).strip()
                        current_section = section_key
                        current_text = []
                        found_section = True
                        break
                
                if not found_section and current_section:
                    current_text.append(line)
            
            # Save any remaining section
            if current_section and current_text:
                sections[current_section] = '\n'.join(current_text).strip()
            
            # If parsing failed, store everything in teaching_lesson
            if not any(sections.values()):
                sections['teaching_lesson'] = response_text
            
            day_plan = DayPlan(
                day_name=day_info['day_name'],
                day_date=day_info['date'],
                learner_outcomes=sections['learner_outcomes'] or 'Content will be generated',
                standards=sections['standards'] or 'Content will be generated',
                materials_needed=sections['materials_needed'] or 'Content will be generated',
                anticipatory_set=sections['anticipatory_set'] or 'Content will be generated',
                teaching_lesson=sections['teaching_lesson'] or response_text,
                modeling=sections['modeling'] or 'Content will be generated',
                instructional_strategies=sections['instructional_strategies'] or 'Content will be generated',
                check_understanding=sections['check_understanding'] or 'Content will be generated',
                guided_practice=sections['guided_practice'] or 'Content will be generated',
                independent_practice=sections['independent_practice'] or 'Content will be generated',
                closure=sections['closure'] or 'Content will be generated',
                summative_assessment=sections['summative_assessment'] or 'Not applicable for today (next major assessment: ' + plan_data.next_major_assessment + ')',
                formative_assessment=sections['formative_assessment'] or 'Content will be generated',
                extended_activities=sections['extended_activities'] or 'Content will be generated',
                review_reteach=sections['review_reteach'] or 'Content will be generated',
                early_finishers=sections['early_finishers'] or 'Content will be generated'
            )
            daily_plans.append(day_plan)
        
        # Create lesson plan
        lesson_plan = LessonPlan(
            user_id=current_user['id'],
            textbook=plan_data.textbook,
            start_date=plan_data.start_date,
            end_date=plan_data.end_date,
            lesson_range=plan_data.lesson_range,
            next_major_assessment=plan_data.next_major_assessment,
            daily_plans=daily_plans
        )
        
        plan_dict = lesson_plan.model_dump()
        plan_dict['created_at'] = plan_dict['created_at'].isoformat()
        
        await db.lesson_plans.insert_one(plan_dict)
        
        return lesson_plan
    except Exception as e:
        logging.error(f"Error creating lesson plan: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating lesson plan: {str(e)}")

@api_router.get("/lesson-plans", response_model=List[LessonPlan])
async def get_lesson_plans(current_user: dict = Depends(get_current_user)):
    plans = await db.lesson_plans.find({"user_id": current_user['id']}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    for plan in plans:
        if isinstance(plan.get('created_at'), str):
            plan['created_at'] = datetime.fromisoformat(plan['created_at'])
    
    return plans

@api_router.get("/lesson-plans/{plan_id}")
async def get_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    plan = await db.lesson_plans.find_one({"id": plan_id, "user_id": current_user['id']}, {"_id": 0})
    if not plan:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    
    if isinstance(plan.get('created_at'), str):
        plan['created_at'] = datetime.fromisoformat(plan['created_at'])
    
    return plan

@api_router.delete("/lesson-plans/{plan_id}")
async def delete_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.lesson_plans.delete_one({"id": plan_id, "user_id": current_user['id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    return {"message": "Lesson plan deleted successfully"}

@api_router.get("/lesson-plans/{plan_id}/export")
async def export_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    plan = await db.lesson_plans.find_one({"id": plan_id, "user_id": current_user['id']}, {"_id": 0})
    if not plan:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('Lesson Plan', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Basic info
    doc.add_paragraph(f"Textbook: {plan['textbook']}")
    doc.add_paragraph(f"Lesson Range: {plan['lesson_range']}")
    doc.add_paragraph(f"Date Range: {plan['start_date']} to {plan['end_date']}")
    doc.add_paragraph(f"Next Major Assessment: {plan['next_major_assessment']}")
    doc.add_paragraph('')
    
    # Daily plans
    for day_plan in plan.get('daily_plans', []):
        doc.add_heading(f"{day_plan['day_name']} - {day_plan['day_date']}", level=1)
        
        sections = [
            ('Learner Outcomes/Objectives', day_plan.get('learner_outcomes', '')),
            ('Standards', day_plan.get('standards', '')),
            ('Materials Needed', day_plan.get('materials_needed', '')),
            ('Anticipatory Set', day_plan.get('anticipatory_set', '')),
            ('Teaching the Lesson', day_plan.get('teaching_lesson', '')),
            ('Modeling', day_plan.get('modeling', '')),
            ('Instructional Strategies', day_plan.get('instructional_strategies', '')),
            ('Check for Understanding', day_plan.get('check_understanding', '')),
            ('Guided Practice/Monitoring', day_plan.get('guided_practice', '')),
            ('Independent Practice', day_plan.get('independent_practice', '')),
            ('Closure', day_plan.get('closure', '')),
            ('Summative Assessment', day_plan.get('summative_assessment', '')),
            ('Formative Assessment', day_plan.get('formative_assessment', '')),
            ('**Extended Activities**', day_plan.get('extended_activities', '')),
            ('**Review and Reteach Activities**', day_plan.get('review_reteach', '')),
            ('**Early Finishers Activities**', day_plan.get('early_finishers', ''))
        ]
        
        for section_title, section_content in sections:
            heading = doc.add_heading(section_title, level=2)
            if '**' in section_title:
                heading.runs[0].bold = True
            doc.add_paragraph(section_content or 'N/A')
        
        doc.add_page_break()
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=lesson_plan_{plan_id}.docx"}
    )


# Class Management Routes
@api_router.post("/classes")
async def create_class(class_data: dict, current_user: dict = Depends(get_current_user)):
    new_class = Class(
        teacher_id=current_user['id'],
        name=class_data['name'],
        description=class_data.get('description')
    )
    class_dict = new_class.model_dump()
    class_dict['created_at'] = class_dict['created_at'].isoformat()
    await db.classes.insert_one(class_dict)
    return new_class

@api_router.get("/classes")
async def get_classes(current_user: dict = Depends(get_current_user)):
    classes = await db.classes.find({"teacher_id": current_user['id']}, {"_id": 0}).to_list(1000)
    
    # Get student counts
    for cls in classes:
        students = await db.students.find({"id": {"$in": cls['student_ids']}}, {"_id": 0, "name": 1, "email": 1}).to_list(1000)
        cls['students'] = students
        cls['student_count'] = len(students)
    
    return classes

# Student Auth Routes
@api_router.post("/auth/student/session")
async def process_student_session(request: Request):
    session_id = request.headers.get('X-Session-ID')
    if not session_id:
        raise HTTPException(status_code=400, detail="Session ID required")
    
    # Call Emergent auth service
    import aiohttp
    async with aiohttp.ClientSession() as session:
        async with session.get(
            'https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data',
            headers={'X-Session-ID': session_id}
        ) as resp:
            if resp.status != 200:
                raise HTTPException(status_code=401, detail="Invalid session")
            user_data = await resp.json()
    
    # Create or get student
    existing_student = await db.students.find_one({"email": user_data['email']}, {"_id": 0})
    if existing_student:
        student = existing_student
    else:
        student = Student(
            name=user_data['name'],
            email=user_data['email'],
            picture=user_data.get('picture')
        )
        student_dict = student.model_dump()
        student_dict['created_at'] = student_dict['created_at'].isoformat()
        await db.students.insert_one(student_dict)
    
    # Create session
    session_token = user_data['session_token']
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    student_session = StudentSession(
        student_id=student['id'],
        session_token=session_token,
        expires_at=expires_at
    )
    session_dict = student_session.model_dump()
    session_dict['expires_at'] = session_dict['expires_at'].isoformat()
    session_dict['created_at'] = session_dict['created_at'].isoformat()
    await db.student_sessions.insert_one(session_dict)
    
    return {
        "student": student,
        "session_token": session_token
    }

async def get_current_student(request: Request):
    # Check cookie first
    session_token = request.cookies.get('student_session_token')
    
    # Fallback to Authorization header
    if not session_token:
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            session_token = auth_header.split(' ')[1]
    
    if not session_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    # Get session
    session = await db.student_sessions.find_one({"session_token": session_token}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=401, detail="Invalid session")
    
    # Check expiry
    expires_at = datetime.fromisoformat(session['expires_at']) if isinstance(session['expires_at'], str) else session['expires_at']
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Session expired")
    
    # Get student
    student = await db.students.find_one({"id": session['student_id']}, {"_id": 0})
    if not student:
        raise HTTPException(status_code=401, detail="Student not found")
    
    return student

@api_router.get("/auth/student/me")
async def get_student_me(request: Request):
    student = await get_current_student(request)
    return student

@api_router.post("/auth/student/logout")
async def student_logout(request: Request):
    session_token = request.cookies.get('student_session_token')
    if session_token:
        await db.student_sessions.delete_one({"session_token": session_token})
    return {"message": "Logged out"}

@api_router.post("/classes/join")
async def join_class(data: dict, request: Request):
    # Get authenticated student
    student = await get_current_student(request)
    
    class_code = data.get('class_code')
    student_id_number = data.get('student_id')
    
    class_data = await db.classes.find_one({"class_code": class_code}, {"_id": 0})
    if not class_data:
        raise HTTPException(status_code=404, detail="Class code not found")
    
    # Update student with student_id if provided
    if student_id_number:
        await db.students.update_one(
            {"id": student['id']},
            {"$set": {"student_id": student_id_number}}
        )
    
    # Add student to class if not already there
    if student['id'] not in class_data['student_ids']:
        await db.classes.update_one(
            {"class_code": class_code},
            {"$push": {"student_ids": student['id']}}
        )
    
    return {"message": "Joined class successfully", "class_name": class_data['name'], "student_id": student['id']}

@api_router.delete("/classes/{class_id}")
async def delete_class(class_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.classes.delete_one({"id": class_id, "teacher_id": current_user['id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Class not found")
    return {"message": "Class deleted successfully"}

# Quiz/Test Routes
@api_router.post("/quizzes/extract-objectives")
async def extract_objectives(data: dict, current_user: dict = Depends(get_current_user)):
    lesson_plan_id = data.get('lesson_plan_id')
    
    # Get lesson plan
    plan = await db.lesson_plans.find_one({"id": lesson_plan_id, "user_id": current_user['id']}, {"_id": 0})
    if not plan:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    
    # Extract objectives and collect all standards
    objectives = []
    all_standards_text = []
    
    for day_plan in plan.get('daily_plans', []):
        # Collect standards from the standards field
        standards_text = day_plan.get('standards', '')
        if standards_text and 'see full plan below' not in standards_text.lower():
            all_standards_text.append(standards_text)
        
        # ALSO check teaching_lesson field for "2. Standards" or "### 2. Standards" section
        teaching_lesson = day_plan.get('teaching_lesson', '')
        if teaching_lesson and '2. Standards' in teaching_lesson:
            # Extract the standards section from teaching_lesson
            import re
            # Look for "### 2. Standards" or "2. Standards" followed by content until next section
            standards_section_pattern = r'(?:###\s*)?2\.\s*Standards?\s*\n(.*?)(?=\n(?:###\s*)?\d+\.|$)'
            match = re.search(standards_section_pattern, teaching_lesson, re.IGNORECASE | re.DOTALL)
            if match:
                all_standards_text.append(match.group(1))
        
        if day_plan.get('learner_outcomes'):
            # Parse objectives (split by line, bullet points, or numbers)
            text = day_plan['learner_outcomes']
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            for line in lines:
                clean_line = line.lstrip('•-*123456789.() ').strip()
                if len(clean_line) > 10:
                    objectives.append({
                        'id': str(uuid.uuid4()),
                        'text': clean_line,
                        'day': day_plan['day_name'],
                        'date': day_plan['day_date'],
                        'selected': True
                    })
    
    # Parse and deduplicate standards - extract only standard numbers/codes
    import re
    unique_standards = set()
    for standards_text in all_standards_text:
        # Skip placeholder text
        if 'see full plan below' in standards_text.lower() or 'content will be generated' in standards_text.lower():
            continue
        
        # Split by lines to process each standard separately
        lines = standards_text.split('\n')
        for line in lines:
            # Look for standard codes at the beginning of lines (before colons or descriptions)
            # Matches: 7.G.1, MS.SS.7.3, CCSS.ELA-LITERACY.RI.RH.6-8.2, etc.
            
            # Pattern 1: Standard code before a colon (e.g., "7.G.1: Description")
            colon_pattern = r'^[\s\-\*\•]*([A-Z0-9][A-Z0-9\.\-]+[0-9A-Z])(?:\s*[:)]|$)'
            match = re.match(colon_pattern, line.strip(), re.IGNORECASE)
            if match:
                code = match.group(1).strip()
                if re.search(r'\d', code) and '.' in code:  # Must have digit and dot
                    unique_standards.add(code)
                continue
            
            # Pattern 2: Standard codes in brackets or bold (e.g., "**7.G.2**")
            bracket_pattern = r'[\[\*]+([A-Z0-9][A-Z0-9\.\-]+[0-9A-Z])[\]\*]+'
            matches = re.findall(bracket_pattern, line, re.IGNORECASE)
            for code in matches:
                if re.search(r'\d', code) and '.' in code:
                    unique_standards.add(code.strip())
    
    # Create standards list
    standards_list = [
        {
            'id': str(uuid.uuid4()),
            'text': std,
            'selected': True
        }
        for std in sorted(unique_standards)
    ]
    
    return {
        "objectives": objectives,
        "standards": standards_list
    }

@api_router.post("/quizzes/generate-questions")
async def generate_questions(data: dict, current_user: dict = Depends(get_current_user)):
    standards_data = data.get('standards', [])  # List of selected standards
    count = data.get('count', 3)  # Number of questions per standard
    
    if not standards_data:
        raise HTTPException(status_code=400, detail="No standards provided")
    
    # Initialize Claude
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    chat = LlmChat(
        api_key=api_key,
        session_id=f"quiz_gen_{current_user['id']}_{datetime.now(timezone.utc).isoformat()}",
        system_message="You are an expert education assessment creator. Generate high-quality multiple choice questions aligned with state educational standards."
    )
    chat.with_model("anthropic", "claude-3-7-sonnet-20250219")
    
    all_questions = []
    
    for standard_code in standards_data:
        prompt = f"""Generate {count} multiple choice questions to assess student understanding of this educational standard:

Standard: {standard_code}

For each question:
1. Make it grade-appropriate and aligned with the standard
2. Provide exactly 4 answer options
3. Indicate which option (0-3) is correct
4. Ensure distractors are plausible but clearly wrong
5. Questions should test knowledge, comprehension, or application related to this standard

Return ONLY a JSON array in this exact format:
[
  {{
    "question_text": "question here",
    "options": ["option 1", "option 2", "option 3", "option 4"],
    "correct_answer": 0,
    "skill": "{standard_code}"
  }}
]

Return ONLY the JSON array, no other text."""

        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        response_text = response if isinstance(response, str) else str(response)
        
        # Parse JSON response
        try:
            # Clean response
            json_start = response_text.find('[')
            json_end = response_text.rfind(']') + 1
            if json_start != -1 and json_end > json_start:
                json_text = response_text[json_start:json_end]
                questions = json.loads(json_text)
                
                for q in questions:
                    all_questions.append({
                        'id': str(uuid.uuid4()),
                        'question_text': q['question_text'],
                        'options': q['options'],
                        'correct_answer': q['correct_answer'],
                        'skill': standard_code
                    })
        except Exception as e:
            logging.error(f"Error parsing questions: {str(e)}")
            continue
    
    return {"questions": all_questions}

@api_router.post("/quizzes")
async def create_quiz(data: dict, current_user: dict = Depends(get_current_user)):
    quiz = QuizTest(
        title=data['title'],
        teacher_id=current_user['id'],
        lesson_plan_id=data['lesson_plan_id'],
        questions=[Question(**q) for q in data['questions']],
        status=data.get('status', 'draft')
    )
    
    quiz_dict = quiz.model_dump()
    quiz_dict['created_at'] = quiz_dict['created_at'].isoformat()
    await db.quizzes.insert_one(quiz_dict)
    
    return quiz

@api_router.get("/quizzes")
async def get_quizzes(current_user: dict = Depends(get_current_user)):
    quizzes = await db.quizzes.find({"teacher_id": current_user['id']}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return quizzes

@api_router.get("/quizzes/{quiz_id}")
async def get_quiz(quiz_id: str):
    quiz = await db.quizzes.find_one({"id": quiz_id}, {"_id": 0})
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    return quiz

@api_router.put("/quizzes/{quiz_id}")
async def update_quiz(quiz_id: str, data: dict, current_user: dict = Depends(get_current_user)):
    result = await db.quizzes.update_one(
        {"id": quiz_id, "teacher_id": current_user['id']},
        {"$set": data}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Quiz not found")
    return {"message": "Quiz updated successfully"}

@api_router.delete("/quizzes/{quiz_id}")
async def delete_quiz(quiz_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.quizzes.delete_one({"id": quiz_id, "teacher_id": current_user['id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Quiz not found")
    return {"message": "Quiz deleted successfully"}

# Assignment Routes
@api_router.post("/assignments")
async def create_assignment(data: dict, current_user: dict = Depends(get_current_user)):
    assignment = Assignment(
        test_id=data['test_id'],
        class_ids=data['class_ids']
    )
    
    assign_dict = assignment.model_dump()
    assign_dict['created_at'] = assign_dict['created_at'].isoformat()
    await db.assignments.insert_one(assign_dict)
    
    # Update quiz status to published
    await db.quizzes.update_one(
        {"id": data['test_id']},
        {"$set": {"status": "published"}}
    )
    
    return assignment

@api_router.get("/assignments/student/{student_id}")
async def get_student_assignments(student_id: str):
    # Get classes the student is in
    classes = await db.classes.find({"student_ids": student_id}, {"_id": 0, "id": 1}).to_list(1000)
    class_ids = [c['id'] for c in classes]
    
    # Get assignments for these classes
    assignments = await db.assignments.find({"class_ids": {"$in": class_ids}}, {"_id": 0}).to_list(1000)
    
    # Get quiz details and check if completed
    result = []
    for assign in assignments:
        quiz = await db.quizzes.find_one({"id": assign['test_id']}, {"_id": 0})
        if quiz:
            submission = await db.submissions.find_one({"test_id": quiz['id'], "student_id": student_id}, {"_id": 0})
            result.append({
                **assign,
                'quiz': quiz,
                'completed': submission is not None,
                'score': submission['score'] if submission else None
            })
    
    return result

# Submission Routes
@api_router.post("/submissions")
async def submit_quiz(data: dict):
    test_id = data['test_id']
    student_id = data['student_id']
    answers = [StudentAnswer(**a) for a in data['answers']]
    
    # Get quiz
    quiz = await db.quizzes.find_one({"id": test_id}, {"_id": 0})
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    # Find the class_id(s) this quiz is assigned to and that this student is in
    assignments = await db.assignments.find({"test_id": test_id}, {"_id": 0}).to_list(1000)
    student_classes = await db.classes.find({"student_ids": student_id}, {"_id": 0, "id": 1}).to_list(1000)
    student_class_ids = [c['id'] for c in student_classes]
    
    # Find matching class
    class_id = None
    for assignment in assignments:
        for assigned_class_id in assignment.get('class_ids', []):
            if assigned_class_id in student_class_ids:
                class_id = assigned_class_id
                break
        if class_id:
            break
    
    if not class_id:
        # Fallback: use first class the student is in
        class_id = student_class_ids[0] if student_class_ids else 'unknown'
    
    # Calculate score and skills breakdown
    correct = 0
    total = len(quiz['questions'])
    skills_breakdown = {}
    
    for answer in answers:
        question = next((q for q in quiz['questions'] if q['id'] == answer.question_id), None)
        if question:
            skill = question['skill']
            if skill not in skills_breakdown:
                skills_breakdown[skill] = {'correct': 0, 'total': 0, 'percentage': 0}
            
            skills_breakdown[skill]['total'] += 1
            
            if question['correct_answer'] == answer.selected_answer:
                correct += 1
                skills_breakdown[skill]['correct'] += 1
    
    # Calculate percentages
    for skill in skills_breakdown:
        skills_breakdown[skill]['percentage'] = (skills_breakdown[skill]['correct'] / skills_breakdown[skill]['total']) * 100
    
    score = (correct / total) * 100 if total > 0 else 0
    
    submission = Submission(
        test_id=test_id,
        student_id=student_id,
        class_id=class_id,
        answers=answers,
        score=score,
        skills_breakdown=skills_breakdown
    )
    
    sub_dict = submission.model_dump()
    sub_dict['submitted_at'] = sub_dict['submitted_at'].isoformat()
    await db.submissions.insert_one(sub_dict)
    
    return submission

# Analytics Routes
@api_router.get("/analytics/class/{class_id}")
async def get_class_analytics(class_id: str, current_user: dict = Depends(get_current_user)):
    # Get all submissions for this class
    submissions = await db.submissions.find({"class_id": class_id}, {"_id": 0}).to_list(10000)
    
    if not submissions:
        return {"message": "No data yet"}
    
    # Get students
    class_data = await db.classes.find_one({"id": class_id}, {"_id": 0})
    students = await db.students.find({"id": {"$in": class_data['student_ids']}}, {"_id": 0}).to_list(1000)
    
    # Calculate analytics
    skill_stats = {}
    student_stats = {}
    
    for sub in submissions:
        # Track student performance
        if sub['student_id'] not in student_stats:
            student = next((s for s in students if s['id'] == sub['student_id']), None)
            student_stats[sub['student_id']] = {
                'student_name': student['name'] if student else 'Unknown',
                'total_score': 0,
                'count': 0,
                'skills': {}
            }
        
        student_stats[sub['student_id']]['total_score'] += sub['score']
        student_stats[sub['student_id']]['count'] += 1
        
        # Track skill performance
        for skill, breakdown in sub['skills_breakdown'].items():
            if skill not in skill_stats:
                skill_stats[skill] = {
                    'skill': skill,
                    'total_attempts': 0,
                    'correct_count': 0,
                    'total_questions': 0,
                    'students_struggling': []
                }
            
            skill_stats[skill]['total_attempts'] += 1
            skill_stats[skill]['correct_count'] += breakdown['correct']
            skill_stats[skill]['total_questions'] += breakdown['total']
            
            # Track student performance on this skill
            if sub['student_id'] not in student_stats[sub['student_id']]['skills']:
                student_stats[sub['student_id']]['skills'][skill] = {'correct': 0, 'total': 0}
            
            student_stats[sub['student_id']]['skills'][skill]['correct'] += breakdown['correct']
            student_stats[sub['student_id']]['skills'][skill]['total'] += breakdown['total']
    
    # Calculate averages and identify struggling students
    for skill, stats in skill_stats.items():
        stats['class_average'] = (stats['correct_count'] / stats['total_questions']) * 100 if stats['total_questions'] > 0 else 0
        
        # Find students struggling with this skill (below 70%)
        for student_id, student_data in student_stats.items():
            if skill in student_data['skills']:
                skill_perf = student_data['skills'][skill]
                percentage = (skill_perf['correct'] / skill_perf['total']) * 100
                if percentage < 70:
                    stats['students_struggling'].append({
                        'student_id': student_id,
                        'student_name': student_data['student_name'],
                        'percentage': percentage
                    })
    
    # Calculate overall student averages
    for student_id, stats in student_stats.items():
        stats['overall_average'] = stats['total_score'] / stats['count'] if stats['count'] > 0 else 0
    
    return {
        'skill_stats': list(skill_stats.values()),
        'student_stats': list(student_stats.values())
    }

@api_router.post("/analytics/remediation-suggestions")
async def get_remediation_suggestions(data: dict, current_user: dict = Depends(get_current_user)):
    skill = data.get('skill')
    student_names = data.get('student_names', [])
    
    # Generate AI suggestions
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    chat = LlmChat(
        api_key=api_key,
        session_id=f"remediation_{current_user['id']}_{datetime.now(timezone.utc).isoformat()}",
        system_message="You are an expert education interventionist providing targeted remediation strategies."
    )
    chat.with_model("anthropic", "claude-3-7-sonnet-20250219")
    
    prompt = f"""Provide exactly 5 specific, actionable remediation activities for students struggling with the following standard/skill:

Standard/Skill: {skill}

Students needing help: {', '.join(student_names) if student_names else 'Multiple students'}

Requirements for EACH of the 5 activities:
- Be specific and immediately actionable in the classroom
- Include materials needed (common classroom items)
- Suggest duration (5-15 minutes)
- Make it engaging and age-appropriate
- Build from concrete to abstract understanding
- Focus on the specific standard/skill listed above

Format: Return exactly 5 activities as a clear numbered list (1-5). Each activity should be 2-3 sentences."""

    user_message = UserMessage(text=prompt)
    response = await chat.send_message(user_message)
    response_text = response if isinstance(response, str) else str(response)
    
    return {"skill": skill, "suggestions": response_text}

# Individual Test Report
@api_router.get("/analytics/test/{quiz_id}")
async def get_test_report(quiz_id: str, current_user: dict = Depends(get_current_user)):
    # Get quiz
    quiz = await db.quizzes.find_one({"id": quiz_id}, {"_id": 0})
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    # Get all submissions for this quiz
    submissions = await db.submissions.find({"test_id": quiz_id}, {"_id": 0}).to_list(10000)
    
    if not submissions:
        return {"message": "No submissions yet"}
    
    # Get student info
    student_ids = [sub['student_id'] for sub in submissions]
    students = await db.students.find({"id": {"$in": student_ids}}, {"_id": 0}).to_list(1000)
    student_map = {s['id']: s for s in students}
    
    # Calculate overall stats
    scores = [sub['score'] for sub in submissions]
    class_average = sum(scores) / len(scores) if scores else 0
    highest_score = max(scores) if scores else 0
    lowest_score = min(scores) if scores else 0
    
    # Calculate standards performance
    standards_stats = {}
    for sub in submissions:
        for standard, breakdown in sub.get('skills_breakdown', {}).items():
            if standard not in standards_stats:
                standards_stats[standard] = {
                    'standard': standard,
                    'total_correct': 0,
                    'total_attempts': 0,
                    'students_struggling': []
                }
            
            standards_stats[standard]['total_correct'] += breakdown['correct']
            standards_stats[standard]['total_attempts'] += breakdown['total']
            
            percentage = (breakdown['correct'] / breakdown['total'] * 100) if breakdown['total'] > 0 else 0
            if percentage < 70:
                student_name = student_map.get(sub['student_id'], {}).get('name', 'Unknown')
                standards_stats[standard]['students_struggling'].append(student_name)
    
    # Calculate standards averages
    standards_list = []
    for std_code, stats in standards_stats.items():
        class_avg = (stats['total_correct'] / stats['total_attempts'] * 100) if stats['total_attempts'] > 0 else 0
        standards_list.append({
            'standard': std_code,
            'class_average': class_avg,
            'students_struggling': len(set(stats['students_struggling']))
        })
    
    # Student results
    student_results = []
    for sub in submissions:
        student_name = student_map.get(sub['student_id'], {}).get('name', 'Unknown')
        
        standards_performance = []
        for standard, breakdown in sub.get('skills_breakdown', {}).items():
            percentage = (breakdown['correct'] / breakdown['total'] * 100) if breakdown['total'] > 0 else 0
            standards_performance.append({
                'standard': standard,
                'percentage': percentage
            })
        
        student_results.append({
            'student_id': sub['student_id'],
            'name': student_name,
            'score': sub['score'],
            'standards_performance': standards_performance
        })
    
    # Sort by score descending
    student_results.sort(key=lambda x: x['score'], reverse=True)
    
    # Question analysis
    question_stats = {}
    for question in quiz['questions']:
        question_stats[question['id']] = {
            'question_text': question['question_text'],
            'standard': question['skill'],
            'correct_count': 0,
            'total_count': 0
        }
    
    for sub in submissions:
        for answer in sub['answers']:
            if answer['question_id'] in question_stats:
                question_stats[answer['question_id']]['total_count'] += 1
                
                # Check if answer is correct
                question = next((q for q in quiz['questions'] if q['id'] == answer['question_id']), None)
                if question and question['correct_answer'] == answer['selected_answer']:
                    question_stats[answer['question_id']]['correct_count'] += 1
    
    questions_analysis = []
    for q_id, stats in question_stats.items():
        percent_correct = (stats['correct_count'] / stats['total_count'] * 100) if stats['total_count'] > 0 else 0
        questions_analysis.append({
            'question_text': stats['question_text'],
            'standard': stats['standard'],
            'percent_correct': percent_correct
        })
    
    return {
        'quiz_title': quiz['title'],
        'total_students': len(submissions),
        'class_average': class_average,
        'highest_score': highest_score,
        'lowest_score': lowest_score,
        'standards': standards_list,
        'students': student_results,
        'questions': questions_analysis
    }

# Individual Student Profile
@api_router.get("/analytics/student/{student_id}")
async def get_student_profile(student_id: str, current_user: dict = Depends(get_current_user)):
    # Get student info
    student = await db.students.find_one({"id": student_id}, {"_id": 0})
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    
    # Get all submissions for this student
    submissions = await db.submissions.find({"student_id": student_id}, {"_id": 0}).sort("submitted_at", 1).to_list(10000)
    
    if not submissions:
        return {
            'student_name': student['name'],
            'tests_taken': 0,
            'overall_average': 0,
            'highest_score': 0,
            'standards_mastered': 0,
            'standards': [],
            'test_history': [],
            'needs_support': []
        }
    
    # Calculate overall stats
    scores = [sub['score'] for sub in submissions]
    overall_average = sum(scores) / len(scores) if scores else 0
    highest_score = max(scores) if scores else 0
    
    # Calculate standards performance
    standards_data = {}
    for sub in submissions:
        for standard, breakdown in sub.get('skills_breakdown', {}).items():
            if standard not in standards_data:
                standards_data[standard] = {
                    'standard': standard,
                    'total_correct': 0,
                    'total_attempts': 0,
                    'scores': []
                }
            
            percentage = (breakdown['correct'] / breakdown['total'] * 100) if breakdown['total'] > 0 else 0
            standards_data[standard]['scores'].append(percentage)
            standards_data[standard]['total_correct'] += breakdown['correct']
            standards_data[standard]['total_attempts'] += breakdown['total']
    
    # Calculate standards averages and mastery
    standards_list = []
    standards_mastered = 0
    needs_support = []
    
    for std_code, data in standards_data.items():
        average = (data['total_correct'] / data['total_attempts'] * 100) if data['total_attempts'] > 0 else 0
        
        if average >= 80:
            standards_mastered += 1
        elif average < 70:
            needs_support.append(std_code)
        
        standards_list.append({
            'standard': std_code,
            'average': average,
            'attempts': len(data['scores'])
        })
    
    # Sort standards by average descending
    standards_list.sort(key=lambda x: x['average'], reverse=True)
    
    # Test history with trends
    test_history = []
    for idx, sub in enumerate(submissions):
        # Get quiz info
        quiz = await db.quizzes.find_one({"id": sub['test_id']}, {"_id": 0})
        
        # Determine trend
        trend = 'stable'
        if idx > 0:
            prev_score = submissions[idx - 1]['score']
            if sub['score'] > prev_score + 5:
                trend = 'up'
            elif sub['score'] < prev_score - 5:
                trend = 'down'
        
        # Standards performance for this test
        standards_performance = []
        for standard, breakdown in sub.get('skills_breakdown', {}).items():
            percentage = (breakdown['correct'] / breakdown['total'] * 100) if breakdown['total'] > 0 else 0
            standards_performance.append({
                'standard': standard,
                'percentage': percentage
            })
        
        test_history.append({
            'quiz_id': sub['test_id'],
            'quiz_title': quiz['title'] if quiz else 'Unknown Quiz',
            'score': sub['score'],
            'date': sub['submitted_at'][:10] if isinstance(sub['submitted_at'], str) else str(sub['submitted_at'])[:10],
            'trend': trend,
            'standards_performance': standards_performance
        })
    
    return {
        'student_name': student['name'],
        'tests_taken': len(submissions),
        'overall_average': overall_average,
        'highest_score': highest_score,
        'standards_mastered': standards_mastered,
        'standards': standards_list,
        'test_history': test_history,
        'needs_support': needs_support
    }

# Admin routes - Invitation Codes
@api_router.post("/admin/invitation-codes")
async def create_invitation_codes(data: CreateInvitationCode, admin_user: dict = Depends(get_admin_user)):
    codes = []
    for _ in range(data.count):
        code = str(uuid.uuid4())[:8].upper()
        invitation = InvitationCode(
            code=code,
            created_by=admin_user['id']
        )
        inv_dict = invitation.model_dump()
        inv_dict['created_at'] = inv_dict['created_at'].isoformat()
        await db.invitation_codes.insert_one(inv_dict)
        codes.append(code)
    
    return {"codes": codes, "count": len(codes)}

@api_router.get("/admin/invitation-codes")
async def get_invitation_codes(admin_user: dict = Depends(get_admin_user)):
    codes = await db.invitation_codes.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Get user info for used codes
    for code in codes:
        if code.get('used_by'):
            user = await db.users.find_one({"id": code['used_by']}, {"_id": 0, "email": 1, "full_name": 1})
            if user:
                code['used_by_email'] = user.get('email')
                code['used_by_name'] = user.get('full_name')
    
    return codes

@api_router.delete("/admin/invitation-codes/{code}")
async def delete_invitation_code(code: str, admin_user: dict = Depends(get_admin_user)):
    result = await db.invitation_codes.delete_one({"code": code})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Invitation code not found")
    return {"message": "Invitation code deleted successfully"}

@api_router.post("/admin/invitation-codes/{code}/deactivate")
async def deactivate_invitation_code(code: str, admin_user: dict = Depends(get_admin_user)):
    result = await db.invitation_codes.update_one({"code": code}, {"$set": {"is_active": False}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Invitation code not found")
    return {"message": "Invitation code deactivated successfully"}

class ChangePassword(BaseModel):
    current_password: str
    new_password: str

# Quiz/Test Models
class Student(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    student_id: Optional[str] = None  # Optional student ID/number
    email: str  # Required for Google OAuth
    picture: Optional[str] = None  # Google profile picture
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StudentSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    student_id: str
    session_token: str
    expires_at: datetime
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Class(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    teacher_id: str
    name: str
    description: Optional[str] = None
    class_code: str = Field(default_factory=lambda: ''.join(__import__('random').choices('ABCDEFGHJKLMNPQRSTUVWXYZ23456789', k=6)))
    student_ids: List[str] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Question(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    question_text: str
    options: List[str]  # 4 options
    correct_answer: int  # index of correct answer (0-3)
    skill: str  # The learning objective/skill being tested

class QuizTest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    teacher_id: str
    lesson_plan_id: str
    questions: List[Question]
    status: str = "draft"  # "draft" or "published"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Assignment(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    test_id: str
    class_ids: List[str]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StudentAnswer(BaseModel):
    question_id: str
    selected_answer: int

class Submission(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    test_id: str
    student_id: str
    class_id: str
    answers: List[StudentAnswer]
    score: float
    skills_breakdown: Dict[str, Dict[str, Any]]  # {skill: {correct: int, total: int, percentage: float}}
    submitted_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

@api_router.post("/auth/change-password")
async def change_password(data: ChangePassword, current_user: dict = Depends(get_current_user)):
    # Verify current password
    if not verify_password(data.current_password, current_user['password']):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    
    # Update password
    new_hashed = hash_password(data.new_password)
    await db.users.update_one(
        {"id": current_user['id']},
        {"$set": {"password": new_hashed}}
    )
    
    return {"message": "Password changed successfully"}

# Admin routes
@api_router.get("/admin/stats")
async def get_admin_stats(admin_user: dict = Depends(get_admin_user)):
    total_users = await db.users.count_documents({"role": "teacher"})
    active_users = await db.users.count_documents({"role": "teacher", "is_active": True})
    inactive_users = total_users - active_users
    total_lesson_plans = await db.lesson_plans.count_documents({})
    
    return AdminStats(
        total_users=total_users,
        active_users=active_users,
        inactive_users=inactive_users,
        total_lesson_plans=total_lesson_plans
    )

@api_router.get("/admin/users")
async def get_all_users(admin_user: dict = Depends(get_admin_user)):
    users = await db.users.find({"role": "teacher"}, {"_id": 0, "password": 0}).to_list(1000)
    
    user_details = []
    for user in users:
        plan_count = await db.lesson_plans.count_documents({"user_id": user['id']})
        user_details.append(UserDetail(
            id=user['id'],
            email=user['email'],
            full_name=user['full_name'],
            state=user.get('state'),
            is_active=user.get('is_active', True),
            join_code=user.get('join_code'),
            created_at=user.get('created_at', ''),
            last_login=user.get('last_login'),
            lesson_plan_count=plan_count
        ))
    
    return user_details

@api_router.post("/admin/users/{user_id}/activate")
async def activate_user(user_id: str, admin_user: dict = Depends(get_admin_user)):
    result = await db.users.update_one({"id": user_id}, {"$set": {"is_active": True}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    return {"message": "User activated successfully"}

@api_router.post("/admin/users/{user_id}/deactivate")
async def deactivate_user(user_id: str, admin_user: dict = Depends(get_admin_user)):
    result = await db.users.update_one({"id": user_id}, {"$set": {"is_active": False}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    return {"message": "User deactivated successfully"}


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
