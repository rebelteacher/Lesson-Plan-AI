"""Lesson plan routes"""
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from datetime import datetime, timezone
from typing import List
import logging
import os
import io

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from emergentintegrations.llm.chat import LlmChat, UserMessage

from models.lesson_plan import LessonPlan, LessonPlanCreate, DayPlan
from utils.database import db
from utils.auth import get_current_user, get_admin_user
from utils.helpers import get_weekdays_between

router = APIRouter(prefix="/lesson-plans", tags=["Lesson Plans"])


@router.post("")
async def create_lesson_plan(plan_data: LessonPlanCreate, current_user: dict = Depends(get_current_user)):
    """Create a new AI-generated lesson plan"""
    try:
        # Get weekdays between start and end date
        weekdays = get_weekdays_between(plan_data.start_date, plan_data.end_date)
        
        if not weekdays:
            raise HTTPException(status_code=400, detail="No weekdays found in the date range")
        
        # Initialize Claude chat
        api_key = os.environ.get('EMERGENT_LLM_KEY')
        
        daily_plans = []
        
        # Generate plan for each day
        for idx, day_info in enumerate(weekdays):
            chat = LlmChat(
                api_key=api_key,
                session_id=f"lesson_plan_{current_user['id']}_{day_info['date']}",
                system_message="You are an expert education consultant helping teachers create detailed daily lesson plans."
            )
            chat.with_model("anthropic", "claude-3-7-sonnet-20250219")
            
            state_standards_text = f"\nState Standards to Align With: {plan_data.state_standards}" if plan_data.state_standards else ""
            
            prompt = f"""Create a detailed lesson plan for {day_info['day_name']}, {day_info['date']} (Day {idx+1} of {len(weekdays)}) based on:

Textbook: {plan_data.textbook}
Lesson Range: {plan_data.lesson_range}
Overall Date Range: {plan_data.start_date} to {plan_data.end_date}
Next Major Assessment: {plan_data.next_major_assessment}{state_standards_text}

Provide specific, actionable content for THIS DAY ONLY for each section:

1. Learner Outcomes/Objectives
2. Standards (include the relevant state standards provided above, formatted clearly)
3. Materials Needed
4. Anticipatory Set
5. Teaching the Lesson
6. Modeling
7. Instructional Strategies
8. Check for Understanding
9. Guided Practice/Monitoring
10. Independent Practice
11. Closure
12. Summative Assessment
13. Formative Assessment
14. Extended Activities
15. Review and Reteach Activities
16. Early Finishers Activities

Make each section detailed and specific to day {idx+1}."""
            
            user_message = UserMessage(text=prompt)
            response = await chat.send_message(user_message)
            response_text = response if isinstance(response, str) else str(response)
            
            # Parse AI response into sections
            sections = parse_lesson_plan_response(response_text)
            
            day_plan = DayPlan(
                day_name=day_info['day_name'],
                day_date=day_info['date'],
                learner_outcomes=sections['learner_outcomes'] or 'Content will be generated',
                standards=sections['standards'] or 'Content will be generated',
                materials_needed=sections['materials_needed'] or 'Content will be generated',
                anticipatory_set=sections['anticipatory_set'] or 'Content will be generated',
                teaching_lesson=sections['teaching_lesson'] or response_text,
                modeling=sections['modeling'] or 'Content will be generated',
                instructional_strategies=sections['instructional_strategies'] or 'Content will be generated',
                check_understanding=sections['check_understanding'] or 'Content will be generated',
                guided_practice=sections['guided_practice'] or 'Content will be generated',
                independent_practice=sections['independent_practice'] or 'Content will be generated',
                closure=sections['closure'] or 'Content will be generated',
                summative_assessment=sections['summative_assessment'] or 'Not applicable for today (next major assessment: ' + plan_data.next_major_assessment + ')',
                formative_assessment=sections['formative_assessment'] or 'Content will be generated',
                extended_activities=sections['extended_activities'] or 'Content will be generated',
                review_reteach=sections['review_reteach'] or 'Content will be generated',
                early_finishers=sections['early_finishers'] or 'Content will be generated'
            )
            daily_plans.append(day_plan)
        
        # Create lesson plan
        lesson_plan = LessonPlan(
            user_id=current_user['id'],
            textbook=plan_data.textbook,
            start_date=plan_data.start_date,
            end_date=plan_data.end_date,
            lesson_range=plan_data.lesson_range,
            next_major_assessment=plan_data.next_major_assessment,
            daily_plans=daily_plans
        )
        
        plan_dict = lesson_plan.model_dump()
        plan_dict['created_at'] = plan_dict['created_at'].isoformat()
        
        await db.lesson_plans.insert_one(plan_dict)
        
        return lesson_plan
    except Exception as e:
        logging.error(f"Error creating lesson plan: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating lesson plan: {str(e)}")


def parse_lesson_plan_response(response_text: str) -> dict:
    """Parse AI response into lesson plan sections"""
    sections = {
        'learner_outcomes': '',
        'standards': '',
        'materials_needed': '',
        'anticipatory_set': '',
        'teaching_lesson': '',
        'modeling': '',
        'instructional_strategies': '',
        'check_understanding': '',
        'guided_practice': '',
        'independent_practice': '',
        'closure': '',
        'summative_assessment': '',
        'formative_assessment': '',
        'extended_activities': '',
        'review_reteach': '',
        'early_finishers': ''
    }
    
    section_mapping = {
        'learner outcomes': 'learner_outcomes',
        'objectives': 'learner_outcomes',
        'standards': 'standards',
        'materials needed': 'materials_needed',
        'anticipatory set': 'anticipatory_set',
        'teaching the lesson': 'teaching_lesson',
        'modeling': 'modeling',
        'instructional strategies': 'instructional_strategies',
        'check for understanding': 'check_understanding',
        'guided practice': 'guided_practice',
        'monitoring': 'guided_practice',
        'independent practice': 'independent_practice',
        'closure': 'closure',
        'summative assessment': 'summative_assessment',
        'formative assessment': 'formative_assessment',
        'extended activities': 'extended_activities',
        'review and reteach': 'review_reteach',
        'reteach activities': 'review_reteach',
        'early finishers': 'early_finishers'
    }
    
    lines = response_text.split('\n')
    current_section = None
    current_text = []
    
    for line in lines:
        line_lower = line.lower().strip()
        found_section = False
        
        for keyword, section_key in section_mapping.items():
            if keyword in line_lower and (line.startswith('#') or line.startswith('**') or line.endswith(':') or len(line) < 50):
                if current_section and current_text:
                    sections[current_section] = '\n'.join(current_text).strip()
                current_section = section_key
                current_text = []
                found_section = True
                break
        
        if not found_section and current_section:
            current_text.append(line)
    
    # Save any remaining section
    if current_section and current_text:
        sections[current_section] = '\n'.join(current_text).strip()
    
    # If parsing failed, store everything in teaching_lesson
    if not any(sections.values()):
        sections['teaching_lesson'] = response_text
    
    return sections


@router.get("", response_model=List[LessonPlan])
async def get_lesson_plans(current_user: dict = Depends(get_current_user)):
    """Get all lesson plans for current user"""
    plans = await db.lesson_plans.find({"user_id": current_user['id']}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    for plan in plans:
        if isinstance(plan.get('created_at'), str):
            plan['created_at'] = datetime.fromisoformat(plan['created_at'])
    
    return plans


@router.get("/{plan_id}")
async def get_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific lesson plan"""
    plan = await db.lesson_plans.find_one({"id": plan_id, "user_id": current_user['id']}, {"_id": 0})
    if not plan:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    
    if isinstance(plan.get('created_at'), str):
        plan['created_at'] = datetime.fromisoformat(plan['created_at'])
    
    return plan


@router.delete("/{plan_id}")
async def delete_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a lesson plan"""
    result = await db.lesson_plans.delete_one({"id": plan_id, "user_id": current_user['id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    return {"message": "Lesson plan deleted successfully"}


@router.get("/{plan_id}/export")
async def export_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    """Export lesson plan to Word document"""
    plan = await db.lesson_plans.find_one({"id": plan_id, "user_id": current_user['id']}, {"_id": 0})
    if not plan:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('Lesson Plan', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Basic info
    doc.add_paragraph(f"Textbook: {plan['textbook']}")
    doc.add_paragraph(f"Lesson Range: {plan['lesson_range']}")
    doc.add_paragraph(f"Date Range: {plan['start_date']} to {plan['end_date']}")
    doc.add_paragraph(f"Next Major Assessment: {plan['next_major_assessment']}")
    doc.add_paragraph('')
    
    # Daily plans
    for day_plan in plan.get('daily_plans', []):
        doc.add_heading(f"{day_plan['day_name']} - {day_plan['day_date']}", level=1)
        
        sections = [
            ('Learner Outcomes/Objectives', day_plan.get('learner_outcomes', '')),
            ('Standards', day_plan.get('standards', '')),
            ('Materials Needed', day_plan.get('materials_needed', '')),
            ('Anticipatory Set', day_plan.get('anticipatory_set', '')),
            ('Teaching the Lesson', day_plan.get('teaching_lesson', '')),
            ('Modeling', day_plan.get('modeling', '')),
            ('Instructional Strategies', day_plan.get('instructional_strategies', '')),
            ('Check for Understanding', day_plan.get('check_understanding', '')),
            ('Guided Practice/Monitoring', day_plan.get('guided_practice', '')),
            ('Independent Practice', day_plan.get('independent_practice', '')),
            ('Closure', day_plan.get('closure', '')),
            ('Summative Assessment', day_plan.get('summative_assessment', '')),
            ('Formative Assessment', day_plan.get('formative_assessment', '')),
            ('Extended Activities', day_plan.get('extended_activities', '')),
            ('Review and Reteach Activities', day_plan.get('review_reteach', '')),
            ('Early Finishers Activities', day_plan.get('early_finishers', ''))
        ]
        
        for section_title, section_content in sections:
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(section_content or 'N/A')
        
        doc.add_page_break()
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=lesson_plan_{plan_id}.docx"}
    )


# Submission Routes
@router.post("/{plan_id}/submit")
async def submit_lesson_plan(plan_id: str, current_user: dict = Depends(get_current_user)):
    """Teacher submits lesson plan to admin for review"""
    plan = await db.lesson_plans.find_one({"id": plan_id, "user_id": current_user['id']})
    if not plan:
        raise HTTPException(status_code=404, detail="Lesson plan not found")
    
    if plan.get('submission_status') == 'pending':
        raise HTTPException(status_code=400, detail="Plan already submitted for review")
    
    # Update submission status
    await db.lesson_plans.update_one(
        {"id": plan_id},
        {"$set": {
            "submission_status": "pending",
            "submitted_at": datetime.now(timezone.utc).isoformat(),
            "admin_feedback": None,
            "reviewed_at": None,
            "reviewed_by": None
        }}
    )
    
    return {"message": "Lesson plan submitted for review"}
